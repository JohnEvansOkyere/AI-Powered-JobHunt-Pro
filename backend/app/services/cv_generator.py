"""
CV Generation Service

Generates tailored CVs for specific job applications using AI.
"""

import json
from typing import Dict, Any, Optional
from datetime import datetime
from sqlalchemy.orm import Session

from app.core.logging import get_logger
from app.core.supabase_client import get_supabase_service_client
from app.services.ai_service import get_ai_service
from app.ai.router import get_model_router
from app.ai.base import TaskType
from app.models.cv import CV
from app.models.job import Job
from app.models.user_profile import UserProfile
from app.models.application import Application
from app.core.config import settings
from app.utils.sanitizer import get_sanitizer

logger = get_logger(__name__)


class CVGenerator:
    """Service for generating tailored CVs for job applications."""
    
    def __init__(self):
        """Initialize CV generator."""
        self.ai_service = get_ai_service()
        self.ai_router = get_model_router()
        self.supabase = get_supabase_service_client()
        self.sanitizer = get_sanitizer()
    
    async def generate_tailored_cv(
        self,
        user_id: str,
        job_id: str,
        db: Session,
        tone: str = "professional",
        highlight_skills: bool = True,
        emphasize_relevant_experience: bool = True
    ) -> Dict[str, Any]:
        """
        Generate a tailored CV for a specific job using the user's original CV as template.

        This function:
        1. Downloads the user's original CV file from storage
        2. Generates tailored text content using AI
        3. Updates the original CV with tailored content (preserves format)
        4. Saves the tailored CV as a new file
        5. Keeps the original CV untouched

        Args:
            user_id: User ID
            job_id: Job ID to tailor CV for
            db: Database session
            tone: Writing tone (professional, confident, friendly)
            highlight_skills: Whether to highlight relevant skills
            emphasize_relevant_experience: Whether to emphasize relevant experience

        Returns:
            dict: Generation result with CV path and metadata
        """
        try:
            # Get user's active CV
            cv = db.query(CV).filter(
                CV.user_id == user_id,
                CV.is_active == True
            ).first()

            if not cv:
                raise ValueError("No active CV found. Please upload a CV first.")

            if not cv.parsed_content:
                raise ValueError("CV has not been parsed yet. Please wait for parsing to complete.")

            # Get job details
            job = db.query(Job).filter(Job.id == job_id).first()
            if not job:
                raise ValueError(f"Job {job_id} not found")

            # Get user profile for additional context
            profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()

            logger.info(f"Generating tailored CV for user {user_id}, job {job_id}")
            logger.info(f"Original CV file: {cv.file_path}, Type: {cv.file_type}")

            # Prepare CV data
            cv_data = cv.parsed_content
            if isinstance(cv_data, str):
                cv_data = json.loads(cv_data)

            # Generate tailored CV content using AI
            tailored_content = await self._generate_cv_content(
                cv_data=cv_data,
                job=job,
                profile=profile,
                tone=tone,
                highlight_skills=highlight_skills,
                emphasize_relevant_experience=emphasize_relevant_experience
            )

            # Download original CV file from storage
            try:
                original_cv_data = self.supabase.storage.from_(settings.SUPABASE_STORAGE_BUCKET).download(cv.file_path)
                logger.info(f"Downloaded original CV from storage: {len(original_cv_data)} bytes")
            except Exception as e:
                logger.error(f"Failed to download original CV: {e}")
                raise ValueError(f"Failed to access original CV file: {str(e)}")

            # Generate tailored CV file (preserving original format)
            file_name, file_bytes, content_type = await self._create_tailored_cv_file(
                original_cv_data=original_cv_data,
                original_file_type=cv.file_type,
                tailored_content=tailored_content,
                cv_data=cv_data,
                job=job
            )

            # Save to Supabase Storage
            storage_path = f"tailored-cvs/{user_id}/{file_name}"
            
            # Upload to Supabase Storage
            try:
                storage_response = self.supabase.storage.from_(settings.SUPABASE_STORAGE_BUCKET).upload(
                    path=storage_path,
                    file=file_bytes,
                    file_options={"content-type": content_type, "upsert": "true"}
                )
                
                # Check for errors (Supabase returns UploadResponse object)
                if hasattr(storage_response, 'error') and storage_response.error:
                    raise ValueError(f"Failed to upload tailored CV: {storage_response.error}")
                elif isinstance(storage_response, dict) and storage_response.get("error"):
                    raise ValueError(f"Failed to upload tailored CV: {storage_response['error']}")
                
                logger.info(f"Tailored CV uploaded successfully to: {storage_path}")
                
            except Exception as storage_error:
                logger.error(f"Supabase Storage upload exception: {storage_error}")
                raise ValueError(f"Failed to upload tailored CV to storage: {str(storage_error)}")
            
            # Get public URL
            try:
                public_url_result = self.supabase.storage.from_(settings.SUPABASE_STORAGE_BUCKET).get_public_url(storage_path)
                public_url = public_url_result if isinstance(public_url_result, str) else public_url_result.get("publicUrl", "")
            except Exception as e:
                logger.warning(f"Could not get public URL: {e}")
                public_url = ""
            
            # Create or update Application record
            application = db.query(Application).filter(
                Application.user_id == user_id,
                Application.job_id == job_id
            ).first()
            
            if application:
                application.tailored_cv_path = storage_path
                application.status = "draft"
                application.generation_settings = {
                    "tone": tone,
                    "highlight_skills": highlight_skills,
                    "emphasize_relevant_experience": emphasize_relevant_experience
                }
                application.updated_at = datetime.utcnow()
            else:
                application = Application(
                    user_id=user_id,
                    job_id=job_id,
                    cv_id=cv.id,
                    tailored_cv_path=storage_path,
                    status="draft",
                    generation_settings={
                        "tone": tone,
                        "highlight_skills": highlight_skills,
                        "emphasize_relevant_experience": emphasize_relevant_experience
                    }
                )
                db.add(application)
            
            db.commit()
            db.refresh(application)
            
            logger.info(f"Successfully generated tailored CV for user {user_id}, job {job_id}")
            
            return {
                "application_id": str(application.id),
                "cv_path": storage_path,
                "public_url": public_url,
                "status": "completed",
                "created_at": application.created_at.isoformat() if application.created_at else None
            }
            
        except Exception as e:
            logger.error(f"Error generating tailored CV: {e}", exc_info=True)
            db.rollback()
            raise
    
    async def _generate_cv_content(
        self,
        cv_data: Dict[str, Any],
        job: Job,
        profile: Optional[UserProfile],
        tone: str,
        highlight_skills: bool,
        emphasize_relevant_experience: bool
    ) -> Dict[str, Any]:
        """
        Generate tailored CV content using AI.

        Args:
            cv_data: Original CV structured data
            job: Job listing
            profile: User profile (optional)
            tone: Writing tone
            highlight_skills: Whether to highlight skills
            emphasize_relevant_experience: Whether to emphasize experience

        Returns:
            dict: Tailored CV content
        """
        # SECURITY: Sanitize all inputs before sending to AI
        logger.info("Sanitizing CV data, job data, and profile data before AI processing")

        # Sanitize CV data
        sanitized_cv = self.sanitizer.sanitize_cv_data(cv_data)

        # Sanitize job data
        job_dict = {
            "title": job.title,
            "company": job.company,
            "location": job.location,
            "job_type": job.job_type,
            "remote_type": job.remote_type,
            "description": job.description,
            "requirements": getattr(job, 'requirements', None),
        }
        sanitized_job = self.sanitizer.sanitize_job_data(job_dict)

        # Sanitize profile data
        profile_dict = None
        if profile:
            profile_dict = {
                "primary_job_title": profile.primary_job_title,
                "seniority_level": profile.seniority_level,
                "work_preference": profile.work_preference,
                "technical_skills": profile.technical_skills,
                "soft_skills": profile.soft_skills,
            }
        sanitized_profile = self.sanitizer.sanitize_profile_data(profile_dict)

        # Check for sensitive data in CV
        cv_text = json.dumps(sanitized_cv)
        sensitive_found = self.sanitizer.check_for_sensitive_data(cv_text)
        if sensitive_found:
            logger.warning(f"Sensitive data types found in CV: {sensitive_found}")

        # Build comprehensive prompt with SANITIZED data
        job_requirements = f"""
Job Title: {sanitized_job.get('title', 'Not specified')}
Company: {sanitized_job.get('company', 'Not specified')}
Location: {sanitized_job.get('location', 'Not specified')}
Job Type: {sanitized_job.get('job_type', 'Not specified')}
Remote Type: {sanitized_job.get('remote_type', 'Not specified')}

Job Description:
{sanitized_job.get('description', 'Not specified')}
"""

        # Add profile context if available
        profile_context = ""
        if sanitized_profile and sanitized_profile.get('primary_job_title'):
            tech_skills = sanitized_profile.get('technical_skills', [])
            soft_skills = sanitized_profile.get('soft_skills', [])
            profile_context = f"""
User Profile Context:
- Primary Job Title: {sanitized_profile.get('primary_job_title', 'Not specified')}
- Seniority Level: {sanitized_profile.get('seniority_level', 'Not specified')}
- Work Preference: {sanitized_profile.get('work_preference', 'Not specified')}
- Technical Skills: {', '.join(tech_skills[:10]) if tech_skills else 'Not specified'}
- Soft Skills: {', '.join(soft_skills[:10]) if soft_skills else 'Not specified'}
"""

        # Convert sanitized CV to JSON (already limited in sanitizer)
        cv_json = json.dumps(sanitized_cv, indent=2)

        prompt = f"""You are an expert career advisor and CV writer. Your task is to tailor this CV for a specific job by REWRITING key sections to emphasize relevant experience and skills.

{job_requirements}

{profile_context}

Original CV Data (Sanitized):
{cv_json}

CRITICAL INSTRUCTIONS - READ CAREFULLY:

1. **DO NOT REMOVE OR OMIT CONTENT**: Include ALL work experience, projects, education, and certifications from the original CV
2. **DO NOT FABRICATE**: Only use factual information from the original CV - never make up dates, companies, or achievements
3. **REWRITE, DON'T REMOVE**: Rewrite bullet points and descriptions to emphasize relevance to this specific job
4. **KEEP ALL METRICS**: Preserve all numbers, percentages, and quantitative achievements (e.g., "90% reduction", "5x faster")
5. **KEEP ALL TECH STACK**: Include all technologies, tools, and frameworks mentioned
6. **REORDER**: Put most relevant experience/skills FIRST, but INCLUDE EVERYTHING
7. **TONE**: Use a {tone} tone throughout

YOUR OUTPUT MUST INCLUDE:
- Professional summary (2-3 sentences, emphasizing fit for THIS job)
- ALL work experience entries (rewritten to highlight relevance)
- ALL projects (rewritten to show how they relate to the job)
- ALL skills (reordered with most relevant first, but include ALL of them)
- ALL education entries
- ALL certifications

Return a JSON object with this COMPLETE structure:
{{
  "summary": "Tailored 2-3 sentence professional summary emphasizing fit for THIS SPECIFIC job role",
  "skills": {{
    "highlighted": ["Top 10-15 skills MOST relevant to this job - reorder from original, don't remove"],
    "all_skills": ["Complete list of ALL skills from original CV - reordered with relevant ones first"]
  }},
  "experience": [
    {{
      "title": "Exact job title from original",
      "company": "Exact company name from original",
      "location": "Exact location from original",
      "start_date": "Exact start date from original",
      "end_date": "Exact end date from original",
      "achievements": [
        "Rewritten achievement 1 emphasizing relevance to target job (KEEP ALL METRICS AND NUMBERS)",
        "Rewritten achievement 2 with quantitative results preserved",
        "Include ALL original bullet points, just reworded for relevance"
      ]
    }}
    // INCLUDE ALL WORK EXPERIENCE FROM ORIGINAL - DO NOT OMIT ANY
  ],
  "projects": [
    {{
      "name": "Exact project name from original",
      "description": "Rewritten description emphasizing how this project relates to the target job",
      "technologies": ["All technologies from original project"],
      "github_url": "GitHub URL if present in original"
    }}
    // INCLUDE ALL PROJECTS FROM ORIGINAL - DO NOT OMIT ANY
  ],
  "education": [
    {{
      "degree": "Exact degree from original",
      "institution": "Exact institution from original",
      "location": "Exact location from original",
      "graduation_date": "Exact date from original"
    }}
    // INCLUDE ALL EDUCATION ENTRIES - DO NOT OMIT ANY
  ],
  "certifications": ["ALL certifications from original CV - exact names"]
}}

VALIDATION CHECKLIST BEFORE RESPONDING:
✓ Does your response include ALL work experience entries from the original?
✓ Does your response include ALL projects from the original?
✓ Does your response include ALL education entries?
✓ Does your response include ALL certifications?
✓ Have you preserved ALL numbers and metrics (e.g., "90%", "5x", "40%")?
✓ Have you preserved ALL technology names and tools?
✓ Did you rewrite (not remove) to emphasize relevance?

Return ONLY valid JSON, no markdown formatting."""

        system_prompt = """You are an expert CV writer who tailors resumes for specific job applications. 
You maintain accuracy while emphasizing relevant experience and skills. You never fabricate information."""
        
        try:
            response = await self.ai_router.generate(
                task_type=TaskType.CV_TAILORING,
                prompt=prompt,
                system_prompt=system_prompt,
                user_id=None,  # Will be passed from caller
                max_tokens=4000,
                temperature=0.7
            )
            
            if not response:
                raise ValueError("AI model returned empty response")
            
            # Parse JSON response
            text = response.strip()
            if text.startswith("```"):
                text = text.split("```")[1]
                if text.startswith("json"):
                    text = text[4:]
            text = text.strip()
            
            tailored_data = json.loads(text)
            return tailored_data
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response as JSON: {e}")
            logger.error(f"Response was: {response[:500]}")
            # Return original CV data if parsing fails
            return cv_data
        except Exception as e:
            logger.error(f"Error generating tailored CV content: {e}")
            # Return original CV data on error
            return cv_data
    
    async def _create_tailored_cv_file(
        self,
        original_cv_data: bytes,
        original_file_type: str,
        tailored_content: Dict[str, Any],
        cv_data: Dict[str, Any],
        job: Job
    ) -> tuple[str, bytes, str]:
        """
        Create a tailored CV file preserving the original format.

        For DOCX: Edits the document directly while preserving formatting
        For PDF: Converts to DOCX, edits, then keeps as DOCX (PDF editing is complex)

        Args:
            original_cv_data: Binary data of original CV file
            original_file_type: File type ('pdf' or 'docx')
            tailored_content: Tailored content from AI
            cv_data: Original parsed CV data
            job: Job listing

        Returns:
            tuple: (filename, file_bytes, content_type)
        """
        import io
        from docx import Document
        from docx.shared import Pt, RGBColor
        from datetime import datetime

        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')

        if original_file_type == 'docx':
            # Edit DOCX directly
            try:
                doc = Document(io.BytesIO(original_cv_data))
                logger.info(f"Loaded DOCX with {len(doc.paragraphs)} paragraphs")

                # Update document with tailored content
                self._update_docx_content(doc, tailored_content, cv_data, job)

                # Save to bytes
                output = io.BytesIO()
                doc.save(output)
                file_bytes = output.getvalue()

                file_name = f"tailored_cv_{job.id}_{timestamp}.docx"
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                logger.info(f"Generated tailored DOCX: {len(file_bytes)} bytes")
                return (file_name, file_bytes, content_type)

            except Exception as e:
                logger.error(f"Error editing DOCX: {e}", exc_info=True)
                # Fallback to creating new document
                return await self._create_new_docx(tailored_content, cv_data, job, timestamp)

        elif original_file_type == 'pdf':
            # For PDF, create a new DOCX with tailored content
            # (Editing PDFs while preserving format is very complex)
            logger.info("PDF detected - creating tailored version as DOCX")
            return await self._create_new_docx(tailored_content, cv_data, job, timestamp)

        else:
            # Unknown format - create new DOCX
            logger.warning(f"Unknown file type: {original_file_type}, creating new DOCX")
            return await self._create_new_docx(tailored_content, cv_data, job, timestamp)

    def _update_docx_content(
        self,
        doc: 'Document',
        tailored_content: Dict[str, Any],
        cv_data: Dict[str, Any],
        job: 'Job'
    ):
        """
        Update DOCX document with tailored content while preserving formatting.

        Strategy:
        1. Find and replace key sections (Professional Summary, etc.)
        2. Keep all formatting, fonts, colors, styles intact
        3. Only update text content
        """
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Update professional summary
        if tailored_content.get("summary"):
            self._replace_text_in_doc(
                doc,
                markers=["professional summary", "summary", "objective", "profile"],
                new_text=tailored_content["summary"],
                section_name="Professional Summary"
            )

        # Update skills section
        if tailored_content.get("skills", {}).get("highlighted"):
            skills_text = ", ".join(tailored_content["skills"]["highlighted"])
            self._replace_text_in_doc(
                doc,
                markers=["skills", "technical skills", "core competencies"],
                new_text=skills_text,
                section_name="Skills"
            )

        # Add a subtle note at the end
        if doc.paragraphs:
            last_para = doc.add_paragraph()
            last_para.add_run(f"\n\nTailored for: {job.title} at {job.company}").font.size = Pt(8)
            last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        logger.info("DOCX content updated with tailored information")

    def _replace_text_in_doc(
        self,
        doc: 'Document',
        markers: list[str],
        new_text: str,
        section_name: str
    ):
        """
        Find and replace text in specific sections of the document.

        Args:
            doc: Document object
            markers: List of possible section headers to look for
            new_text: New text to insert
            section_name: Name of section (for logging)
        """
        found = False
        for i, para in enumerate(doc.paragraphs):
            para_text_lower = para.text.lower().strip()

            # Check if this paragraph is a section header
            for marker in markers:
                if marker in para_text_lower and len(para_text_lower) < 50:
                    # Found the section header
                    # Replace the NEXT paragraph(s) with new content
                    if i + 1 < len(doc.paragraphs):
                        next_para = doc.paragraphs[i + 1]

                        # Preserve formatting of first run
                        if next_para.runs:
                            run = next_para.runs[0]
                            font_name = run.font.name
                            font_size = run.font.size
                            font_bold = run.font.bold
                            font_italic = run.font.italic

                            # Clear and update text
                            next_para.clear()
                            new_run = next_para.add_run(new_text)

                            # Restore formatting
                            if font_name:
                                new_run.font.name = font_name
                            if font_size:
                                new_run.font.size = font_size
                            if font_bold:
                                new_run.font.bold = font_bold
                            if font_italic:
                                new_run.font.italic = font_italic

                            logger.info(f"Updated {section_name} section")
                            found = True
                            return

        if not found:
            logger.warning(f"Could not find {section_name} section in document")

    async def _create_new_docx(
        self,
        tailored_content: Dict[str, Any],
        cv_data: Dict[str, Any],
        job: 'Job',
        timestamp: str
    ) -> tuple[str, bytes, str]:
        """
        Create a new DOCX document from scratch with tailored content.

        Used as fallback when original format can't be preserved.

        Returns:
            tuple: (filename, file_bytes, content_type)
        """
        import io
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # Header with personal info
        personal_info = cv_data.get("personal_info", {})
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal_info.get("name", "Your Name"))
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Contact info
        contact_para = doc.add_paragraph()
        contact_parts = []
        if personal_info.get("email"):
            contact_parts.append(personal_info["email"])
        if personal_info.get("phone"):
            contact_parts.append(personal_info["phone"])
        if personal_info.get("location"):
            contact_parts.append(personal_info["location"])

        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()  # Spacer

        # Professional Summary
        if tailored_content.get("summary"):
            summary_heading = doc.add_paragraph()
            summary_heading.add_run("PROFESSIONAL SUMMARY").font.bold = True
            summary_heading.add_run().font.size = Pt(14)

            summary_para = doc.add_paragraph(tailored_content["summary"])
            summary_para.paragraph_format.space_after = Pt(12)

        # Key Skills
        if tailored_content.get("skills", {}).get("highlighted"):
            skills_heading = doc.add_paragraph()
            skills_heading.add_run("KEY SKILLS").font.bold = True

            skills_text = ", ".join(tailored_content["skills"]["highlighted"])
            skills_para = doc.add_paragraph(skills_text)
            skills_para.paragraph_format.space_after = Pt(12)

        # Experience
        if tailored_content.get("experience"):
            exp_heading = doc.add_paragraph()
            exp_heading.add_run("PROFESSIONAL EXPERIENCE").font.bold = True
            exp_heading.paragraph_format.space_after = Pt(6)

            for exp in tailored_content["experience"]:
                # Job title and company
                job_para = doc.add_paragraph()
                title_text = f"{exp.get('title', '')} - {exp.get('company', '')}"
                if exp.get('location'):
                    title_text += f" ({exp['location']})"
                job_run = job_para.add_run(title_text)
                job_run.font.bold = True
                job_run.font.size = Pt(11)

                # Dates
                if exp.get("start_date") or exp.get("end_date"):
                    dates_para = doc.add_paragraph()
                    dates_run = dates_para.add_run(f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}")
                    dates_run.font.italic = True
                    dates_run.font.size = Pt(10)

                # Achievements (bullet points)
                if exp.get("achievements"):
                    for achievement in exp["achievements"]:
                        bullet_para = doc.add_paragraph(style='List Bullet')
                        bullet_para.add_run(achievement).font.size = Pt(10)

                doc.add_paragraph()  # Spacer

        # Projects
        if tailored_content.get("projects"):
            projects_heading = doc.add_paragraph()
            projects_heading.add_run("PROJECTS").font.bold = True
            projects_heading.paragraph_format.space_after = Pt(6)

            for proj in tailored_content["projects"]:
                # Project name
                proj_para = doc.add_paragraph()
                proj_run = proj_para.add_run(proj.get('name', ''))
                proj_run.font.bold = True
                proj_run.font.size = Pt(11)

                # GitHub URL
                if proj.get('github_url'):
                    doc.add_paragraph(f"GitHub: {proj['github_url']}").font.size = Pt(9)

                # Description
                if proj.get('description'):
                    desc_para = doc.add_paragraph(proj['description'])
                    desc_para.font.size = Pt(10)

                # Technologies
                if proj.get('technologies'):
                    tech_para = doc.add_paragraph()
                    tech_para.add_run("Technologies: ").font.bold = True
                    tech_para.add_run(", ".join(proj['technologies'])).font.size = Pt(10)

                doc.add_paragraph()  # Spacer

        # Education
        if tailored_content.get("education"):
            edu_heading = doc.add_paragraph()
            edu_heading.add_run("EDUCATION").font.bold = True
            edu_heading.paragraph_format.space_after = Pt(6)

            for edu in tailored_content["education"]:
                # Degree and institution
                edu_para = doc.add_paragraph()
                edu_text = f"{edu.get('degree', '')} - {edu.get('institution', '')}"
                if edu.get('location'):
                    edu_text += f" ({edu['location']})"
                edu_run = edu_para.add_run(edu_text)
                edu_run.font.bold = True
                edu_run.font.size = Pt(11)

                # Graduation date
                if edu.get('graduation_date'):
                    date_para = doc.add_paragraph()
                    date_para.add_run(edu['graduation_date']).font.italic = True

                doc.add_paragraph()  # Spacer

        # Certifications
        if tailored_content.get("certifications"):
            cert_heading = doc.add_paragraph()
            cert_heading.add_run("CERTIFICATIONS").font.bold = True
            cert_heading.paragraph_format.space_after = Pt(6)

            for cert in tailored_content["certifications"]:
                cert_para = doc.add_paragraph(style='List Bullet')
                cert_para.add_run(cert).font.size = Pt(10)

        # Footer note
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(f"\nTailored for: {job.title} at {job.company}")
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        file_bytes = output.getvalue()

        file_name = f"tailored_cv_{job.id}_{timestamp}.docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        logger.info(f"Created new tailored DOCX: {len(file_bytes)} bytes")
        return (file_name, file_bytes, content_type)

    def _format_cv_document(
        self,
        tailored_content: Dict[str, Any],
        original_cv_data: Dict[str, Any],
        job: Job
    ) -> str:
        """
        Format tailored CV content as a Markdown document.
        
        Args:
            tailored_content: Tailored CV content from AI
            original_cv_data: Original CV data (for personal info)
            job: Job listing
            
        Returns:
            str: Formatted CV document (Markdown)
        """
        lines = []
        
        # Header with personal info
        personal_info = original_cv_data.get("personal_info", {})
        lines.append("# " + personal_info.get("name", "Your Name"))
        lines.append("")
        if personal_info.get("email"):
            lines.append(f"Email: {personal_info['email']}")
        if personal_info.get("phone"):
            lines.append(f"Phone: {personal_info['phone']}")
        if personal_info.get("location"):
            lines.append(f"Location: {personal_info['location']}")
        if personal_info.get("linkedin"):
            lines.append(f"LinkedIn: {personal_info['linkedin']}")
        if personal_info.get("github"):
            lines.append(f"GitHub: {personal_info['github']}")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Professional Summary
        if tailored_content.get("summary"):
            lines.append("## Professional Summary")
            lines.append("")
            lines.append(tailored_content["summary"])
            lines.append("")
        
        # Highlighted Skills
        if tailored_content.get("skills", {}).get("highlighted"):
            lines.append("## Key Skills")
            lines.append("")
            skills = tailored_content["skills"]["highlighted"]
            lines.append(", ".join(skills))
            lines.append("")
        
        # Experience
        if tailored_content.get("experience"):
            lines.append("## Professional Experience")
            lines.append("")
            for exp in tailored_content["experience"]:
                lines.append(f"### {exp.get('title', '')} at {exp.get('company', '')}")
                if exp.get("location"):
                    lines.append(f"*{exp.get('location')}*")
                if exp.get("start_date") or exp.get("end_date"):
                    dates = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
                    lines.append(f"*{dates}*")
                lines.append("")
                if exp.get("description"):
                    lines.append(exp["description"])
                    lines.append("")
                if exp.get("achievements"):
                    lines.append("**Key Achievements:**")
                    for achievement in exp["achievements"]:
                        lines.append(f"- {achievement}")
                    lines.append("")
                if exp.get("relevance_note"):
                    lines.append(f"*Note: {exp['relevance_note']}*")
                    lines.append("")
        
        # Education
        if tailored_content.get("education"):
            lines.append("## Education")
            lines.append("")
            for edu in tailored_content["education"]:
                lines.append(f"### {edu.get('degree', '')}")
                lines.append(f"{edu.get('institution', '')}")
                if edu.get("location"):
                    lines.append(f"*{edu.get('location')}*")
                if edu.get("graduation_date"):
                    lines.append(f"*Graduated: {edu.get('graduation_date')}*")
                if edu.get("gpa"):
                    lines.append(f"*GPA: {edu.get('gpa')}*")
                lines.append("")
        
        # Skills (Full List)
        skills_data = tailored_content.get("skills", {})
        if skills_data.get("technical"):
            lines.append("## Technical Skills")
            lines.append("")
            lines.append(", ".join(skills_data["technical"]))
            lines.append("")
        
        # Projects
        if tailored_content.get("projects"):
            lines.append("## Projects")
            lines.append("")
            for project in tailored_content["projects"]:
                lines.append(f"### {project.get('name', '')}")
                if project.get("description"):
                    lines.append(project["description"])
                if project.get("technologies"):
                    lines.append(f"**Technologies:** {', '.join(project['technologies'])}")
                if project.get("relevance_note"):
                    lines.append(f"*{project['relevance_note']}*")
                lines.append("")
        
        # Footer note
        lines.append("---")
        lines.append("")
        lines.append(f"*Tailored for: {job.title} at {job.company}*")
        lines.append(f"*Generated on: {datetime.utcnow().strftime('%B %d, %Y')}*")
        
        return "\n".join(lines)


# Global service instance
_cv_generator: Optional[CVGenerator] = None


def get_cv_generator() -> CVGenerator:
    """Get or create the global CV generator instance."""
    global _cv_generator
    if _cv_generator is None:
        _cv_generator = CVGenerator()
    return _cv_generator

