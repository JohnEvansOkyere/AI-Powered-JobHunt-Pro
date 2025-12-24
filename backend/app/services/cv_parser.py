"""
CV Parsing Service

Extracts structured data from CV files (PDF/DOCX) using libraries and AI.
"""

import io
import json
from typing import Dict, Any, Optional
from pathlib import Path

from app.core.logging import get_logger
from app.ai.router import get_model_router
from app.ai.base import TaskType

logger = get_logger(__name__)


class CVParser:
    """Service for parsing CV files and extracting structured data."""

    def __init__(self):
        self.ai_router = get_model_router()

    def parse_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file bytes
            
        Returns:
            Extracted text content
        """
        try:
            import PyPDF2
            
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            # Fallback to pdfplumber if PyPDF2 fails
            try:
                import pdfplumber
                
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        text_parts.append(page.extract_text() or "")
                    return "\n\n".join(text_parts)
            except Exception as e2:
                logger.error(f"Error with pdfplumber fallback: {e2}")
                raise ValueError(f"Failed to parse PDF: {str(e)}")

    def parse_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file bytes
            
        Returns:
            Extracted text content
        """
        try:
            from docx import Document
            
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """
        Extract raw text from CV file.
        
        Args:
            file_content: File bytes
            file_type: File type ('pdf' or 'docx')
            
        Returns:
            Extracted text
        """
        if file_type.lower() == 'pdf':
            return self.parse_pdf(file_content)
        elif file_type.lower() in ['docx', 'doc']:
            return self.parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    async def extract_structured_data(self, raw_text: str) -> Dict[str, Any]:
        """
        Use AI to extract structured data from CV text.
        
        Args:
            raw_text: Raw text extracted from CV
            
        Returns:
            Structured data dictionary
        """
        prompt = f"""Extract structured information from this CV/resume. Return a JSON object with the following structure:

{{
  "personal_info": {{
    "name": "Full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, Country",
    "linkedin": "LinkedIn URL if present",
    "github": "GitHub URL if present",
    "website": "Personal website if present"
  }},
  "summary": "Professional summary or objective",
  "experience": [
    {{
      "title": "Job title",
      "company": "Company name",
      "location": "Location",
      "start_date": "Start date",
      "end_date": "End date or 'Present'",
      "description": "Job description",
      "achievements": ["Achievement 1", "Achievement 2"]
    }}
  ],
  "education": [
    {{
      "degree": "Degree name",
      "institution": "Institution name",
      "location": "Location",
      "graduation_date": "Graduation date",
      "gpa": "GPA if mentioned"
    }}
  ],
  "skills": {{
    "technical": ["Skill 1", "Skill 2"],
    "languages": ["Language 1", "Language 2"],
    "certifications": ["Certification 1", "Certification 2"]
  }},
  "projects": [
    {{
      "name": "Project name",
      "description": "Project description",
      "technologies": ["Tech 1", "Tech 2"],
      "url": "Project URL if present"
    }}
  ]
}}

CV Text:
{raw_text[:4000]}  # Limit to avoid token limits

Return only valid JSON, no markdown formatting."""

        try:
            response = await self.ai_router.generate(
                task_type=TaskType.CV_PARSING,  # Use CV parsing task type
                prompt=prompt,
                max_tokens=2000,
            )
            
            if not response:
                raise ValueError("AI model returned empty response")
            
            # Parse JSON response
            # Remove markdown code blocks if present
            text = response.strip()
            if text.startswith("```"):
                text = text.split("```")[1]
                if text.startswith("json"):
                    text = text[4:]
            text = text.strip()
            
            structured_data = json.loads(text)
            return structured_data
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response as JSON: {e}")
            logger.error(f"Response was: {response[:500]}")
            # Return basic structure if parsing fails
            return {
                "personal_info": {},
                "summary": "",
                "experience": [],
                "education": [],
                "skills": {"technical": [], "languages": [], "certifications": []},
                "projects": [],
            }
        except Exception as e:
            logger.error(f"Error extracting structured data with AI: {e}")
            # Return basic structure on error
            return {
                "personal_info": {},
                "summary": "",
                "experience": [],
                "education": [],
                "skills": {"technical": [], "languages": [], "certifications": []},
                "projects": [],
            }

    async def parse_cv(
        self, file_content: bytes, file_type: str
    ) -> Dict[str, Any]:
        """
        Complete CV parsing pipeline: extract text and structure.
        
        Args:
            file_content: CV file bytes
            file_type: File type ('pdf' or 'docx')
            
        Returns:
            Dictionary with 'raw_text' and 'parsed_content'
        """
        try:
            # Extract raw text
            raw_text = self.extract_text(file_content, file_type)
            
            if not raw_text or len(raw_text.strip()) < 50:
                raise ValueError("CV appears to be empty or unreadable")
            
            # Extract structured data using AI
            parsed_content = await self.extract_structured_data(raw_text)
            
            return {
                "raw_text": raw_text,
                "parsed_content": parsed_content,
            }
        except Exception as e:
            logger.error(f"Error parsing CV: {e}")
            raise

